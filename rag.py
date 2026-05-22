import faiss
import numpy as np
import requests
import io
import os
import pickle
import zipfile
import hashlib
import shelve
import threading
from concurrent.futures import ThreadPoolExecutor, as_completed
from dotenv import load_dotenv
load_dotenv()

# -----------------------
# Config from .env
# -----------------------
OLLAMA_BASE_URL = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")
EMBED_MODEL     = os.getenv("EMBED_MODEL",     "nomic-embed-text")
GENERATE_MODEL  = os.getenv("GENERATE_MODEL",  "qwen2:1.5b")
TOP_K           = int(os.getenv("RETRIEVE_TOP_K", 8))
CHUNK_SIZE      = int(os.getenv("CHUNK_SIZE",     5000))
CHUNK_OVERLAP   = int(os.getenv("CHUNK_OVERLAP",  300))
TEMPERATURE     = float(os.getenv("TEMPERATURE",  0.3))
EMBED_WORKERS   = int(os.getenv("EMBED_WORKERS",  8))
CACHE_PATH      = os.getenv("EMBED_CACHE_PATH",   "embed_cache")


cache_lock = threading.Lock()
# -----------------------
# Load TXT
# -----------------------
def load_txt(file):
    text = file.read().decode("utf-8")
    return [line.strip() for line in text.split("\n") if line.strip()]

# -----------------------
# Load PDF — handles multi-column layouts
# -----------------------
def load_pdf(file):
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("Install pdfplumber: pip install pdfplumber")

    lines = []
    try:
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                page_width = page.width

                # Split page into left and right columns
                left_bbox  = (0,             0, page_width / 2, page.height)
                right_bbox = (page_width / 2, 0, page_width,    page.height)

                left_text  = page.within_bbox(left_bbox).extract_text()
                right_text = page.within_bbox(right_bbox).extract_text()

                # If both halves have text treat as two-column
                # otherwise fall back to normal single column extraction
                if left_text and right_text:
                    combined = (left_text or "") + "\n" + (right_text or "")
                else:
                    combined = page.extract_text() or ""

                for line in combined.split("\n"):
                    line = line.strip()
                    # Skip very short lines — page numbers, headers etc
                    if line and len(line) > 3:
                        lines.append(line)

    except Exception as e:
        raise ValueError(f"Could not read PDF '{file.name}': {e}")
    return lines

# -----------------------
# Load DOCX — with raw XML fallback
# -----------------------
def load_docx(file):
    raw = file.read()
    lines = []

    # First attempt: parse properly with python-docx
    try:
        from docx import Document
        doc = Document(io.BytesIO(raw))
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                lines.append(text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        lines.append(text)
        if lines:
            return lines
    except Exception:
        pass

    # Second attempt: open as zip and extract raw XML text
    try:
        import re
        with zipfile.ZipFile(io.BytesIO(raw)) as z:
            with z.open("word/document.xml") as xml_file:
                xml_content = xml_file.read().decode("utf-8", errors="ignore")
        text = re.sub(r"<[^>]+>", " ", xml_content)
        for line in text.split():
            line = line.strip()
            if line:
                lines.append(line)
        if lines:
            return lines
    except Exception:
        pass

    # Third attempt: plain text decode as last resort
    try:
        text = raw.decode("utf-8", errors="ignore")
        lines = [line.strip() for line in text.split("\n") if line.strip()]
        if lines:
            return lines
    except Exception:
        pass

    raise ValueError(
        f"Could not extract text from '{file.name}'. "
        "The file may be severely corrupted. Please try re-saving it."
    )

# -----------------------
# Chunking
# -----------------------
def chunk_text(lines):
    full_text = " ".join(lines)
    chunks = []
    start = 0
    while start < len(full_text):
        end = start + CHUNK_SIZE
        chunk = full_text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        start += CHUNK_SIZE - CHUNK_OVERLAP
    return chunks

# -----------------------
# Universal file loader
# -----------------------
def load_file(file):
    name = file.name.lower()
    if name.endswith(".txt"):
        lines = load_txt(file)
    elif name.endswith(".pdf"):
        lines = load_pdf(file)
    elif name.endswith(".docx"):
        lines = load_docx(file)
    else:
        try:
            text = file.read().decode("utf-8", errors="ignore")
            lines = [line.strip() for line in text.split("\n") if line.strip()]
        except Exception:
            raise ValueError(f"Unsupported file type: {file.name}")
    return chunk_text(lines)

# -----------------------
# Embedding with disk cache
# -----------------------
def embed(text):
    key = hashlib.md5(text.encode()).hexdigest()

    # Check cache
    with cache_lock:
        with shelve.open(CACHE_PATH) as cache:
            if key in cache:
                return cache[key]

    try:
        res = requests.post(
            f"{OLLAMA_BASE_URL}/api/embeddings",
            json={
                "model": EMBED_MODEL,
                "prompt": text
            },
            timeout=120
        )

        # Show actual Ollama error
        if res.status_code != 200:
            print("OLLAMA ERROR:")
            print(res.text)

            raise Exception(
                f"Embedding failed with status {res.status_code}"
            )

        result = res.json()["embedding"]

        # Save safely
        with cache_lock:
            with shelve.open(CACHE_PATH) as cache:
                cache[key] = result

        return result

    except Exception as e:
        raise Exception(f"Embedding Error: {e}")
# -----------------------
# Build FAISS index (parallel embedding)
# -----------------------
def build_index(docs, progress_callback=None):
    vectors = [None] * len(docs)

    with ThreadPoolExecutor(max_workers=EMBED_WORKERS) as executor:
        futures = {executor.submit(embed, doc): i for i, doc in enumerate(docs)}
        completed = 0
        for future in as_completed(futures):
            i = futures[future]
            vectors[i] = future.result()
            completed += 1
            if progress_callback:
                progress_callback(completed, len(docs))

    vectors = np.array(vectors).astype("float32")
    dim = vectors.shape[1]
    index = faiss.IndexFlatL2(dim)
    index.add(vectors)
    return index

# -----------------------
# Save index to disk
# -----------------------
def save_index(index, docs, path="saved_index.pkl"):
    with open(path, "wb") as f:
        pickle.dump({
            "index": faiss.serialize_index(index),
            "docs": docs
        }, f)

# -----------------------
# Load index from disk
# -----------------------
def load_saved_index(path="saved_index.pkl"):
    if not os.path.exists(path):
        return None, None
    with open(path, "rb") as f:
        data = pickle.load(f)
    return faiss.deserialize_index(data["index"]), data["docs"]

# -----------------------
# Retrieve
# -----------------------
def retrieve(query, index, docs, k=TOP_K):
    q_vec = np.array([embed(query)]).astype("float32")
    _, idx = index.search(q_vec, k)
    return [docs[i].strip() for i in idx[0]]

# -----------------------
# Generate Answer
# -----------------------
def generate(context, question):
    context_text = "\n".join(context)

    prompt = f"""
You are a helpful assistant that answers questions based on the provided context.

Instructions:
1. Answer the question using the information from the context below.
2. If the context partially covers the question, answer what you can from it.
3. Be detailed and complete in your answer.
4. Only say "Not found in context" if the context has absolutely zero relevant information.

Context:
{context_text}

Question:
{question}

Answer:
"""

    res = requests.post(
        f"{OLLAMA_BASE_URL}/api/generate",
        json={
            "model": GENERATE_MODEL,
            "prompt": prompt,
            "stream": False,
            "options": {
                "temperature": TEMPERATURE
            }
        }
    )

    return res.json().get("response", "No response")   